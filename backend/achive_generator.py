"""员工档案生成器"""
"""
档案生成模块
根据员工信息生成 Word 格式的人事档案
"""

from pathlib import Path
from datetime import datetime
from typing import Dict, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ 警告: python-docx 未安装，请运行: pip install python-docx")


class ArchiveGenerator:
    """档案生成器"""
    
    def __init__(self, template_path: Optional[str] = None):
        self.template_path = template_path
    
    def generate_word(self, employee_info: Dict, output_path: str) -> bool:
        """
        生成 Word 格式的员工档案
        
        Args:
            employee_info: 员工信息字典（包含姓名、身份证号等）
            output_path: 输出文件路径
            
        Returns:
            bool: 是否生成成功
        """
        if not DOCX_AVAILABLE:
            return self._generate_text_file(employee_info, output_path)
        
        try:
            doc = Document()
            
            # 设置文档标题
            title = doc.add_heading(f"{employee_info.get('name', '员工')} 人事档案", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # ========== 1. 基本信息 ==========
            doc.add_heading("一、基本信息", level=1)
            table = doc.add_table(rows=8, cols=2)
            table.style = 'Light Grid Accent 1'
            
            basic_fields = [
                ("姓名", employee_info.get('name', '')),
                ("性别", employee_info.get('gender', '')),
                ("民族", employee_info.get('ethnicity', '')),
                ("身份证号", employee_info.get('id_card', '')),
                ("出生日期", employee_info.get('birth_date', '')),
                ("联系电话", employee_info.get('phone', '')),
                ("电子邮箱", employee_info.get('email', '')),
                ("政治面貌", employee_info.get('political_status', '')),
            ]
            
            for i, (label, value) in enumerate(basic_fields):
                table.cell(i, 0).text = label
                table.cell(i, 1).text = str(value) if value else '待补充'
            
            # ========== 2. 住址信息 ==========
            doc.add_heading("二、住址信息", level=1)
            addr_table = doc.add_table(rows=2, cols=2)
            addr_table.style = 'Light Grid Accent 1'
            
            addr_fields = [
                ("户籍地址", employee_info.get('permanent_address', '')),
                ("现住址", employee_info.get('current_address', '')),
            ]
            
            for i, (label, value) in enumerate(addr_fields):
                addr_table.cell(i, 0).text = label
                addr_table.cell(i, 1).text = str(value) if value else '待补充'
            
            # ========== 3. 教育信息 ==========
            doc.add_heading("三、教育信息", level=1)
            edu_table = doc.add_table(rows=5, cols=2)
            edu_table.style = 'Light Grid Accent 1'
            
            edu_fields = [
                ("最高学历", employee_info.get('education', '')),
                ("毕业院校", employee_info.get('school', '')),
                ("专业", employee_info.get('major', '')),
                ("毕业时间", employee_info.get('graduation_date', '')),
                ("学习方式", employee_info.get('mode_of_study', '')),
            ]
            
            for i, (label, value) in enumerate(edu_fields):
                edu_table.cell(i, 0).text = label
                edu_table.cell(i, 1).text = str(value) if value else '待补充'
            
            # ========== 4. 工作经历 ==========
            work_experiences = employee_info.get('work_experiences', [])
            if work_experiences:
                doc.add_heading("四、工作经历", level=1)
                for exp in work_experiences:
                    company = exp.get('company_name', '未知公司')
                    position = exp.get('position', '未知职位')
                    start_date = exp.get('start_date', '')
                    end_date = exp.get('end_date', '至今')
                    
                    p = doc.add_paragraph()
                    p.add_run(f"{company}").bold = True
                    p.add_run(f" ｜ {position}")
                    doc.add_paragraph(f"  时间：{start_date} - {end_date}", style='List Bullet')
                    if exp.get('job_description'):
                        doc.add_paragraph(f"  职责：{exp.get('job_description')}", style='List Bullet')
            else:
                doc.add_heading("四、工作经历", level=1)
                doc.add_paragraph("无工作经历")
            
            # ========== 5. 任职信息 ==========
            doc.add_heading("五、任职信息", level=1)
            job_table = doc.add_table(rows=6, cols=2)
            job_table.style = 'Light Grid Accent 1'
            
            job_fields = [
                ("部门", employee_info.get('department', '')),
                ("职位", employee_info.get('position', '')),
                ("入职日期", employee_info.get('hire_date', '')),
                ("职级", employee_info.get('job_level', '')),
                ("合同类型", employee_info.get('contract_type', '')),
                ("合同期限", employee_info.get('contract_duration', '')),
            ]
            
            for i, (label, value) in enumerate(job_fields):
                job_table.cell(i, 0).text = label
                job_table.cell(i, 1).text = str(value) if value else '待补充'
            
            # ========== 6. 资格证书 ==========
            certificates = employee_info.get('professional_qualifications', [])
            if certificates:
                doc.add_heading("六、专业资格证书", level=1)
                for cert in certificates:
                    doc.add_paragraph(
                        f"• {cert.get('certificate_name', '未知证书')}",
                        style='List Bullet'
                    )
            
            # ========== 7. 荣誉奖励 ==========
            honors = employee_info.get('honors_and_awards', [])
            if honors:
                doc.add_heading("七、荣誉与奖励", level=1)
                for honor in honors:
                    doc.add_paragraph(
                        f"• {honor.get('title', '未知奖项')}",
                        style='List Bullet'
                    )
            
            # ========== 8. 归档信息 ==========
            doc.add_heading("八、归档信息", level=1)
            doc.add_paragraph(f"档案生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"档案状态：{employee_info.get('status', '已生成')}")
            
            # 保存文档
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"生成 Word 档案失败: {e}")
            return self._generate_text_file(employee_info, output_path)
    
    def _generate_text_file(self, employee_info: Dict, output_path: str) -> bool:
        """生成纯文本格式的档案（备用方案）"""
        try:
            content = f"""
{'='*40}
          {employee_info.get('name', '员工')} 人事档案
{'='*40}

【一、基本信息】
姓名：{employee_info.get('name', '待补充')}
性别：{employee_info.get('gender', '待补充')}
民族：{employee_info.get('ethnicity', '待补充')}
身份证号：{employee_info.get('id_card', '待补充')}
出生日期：{employee_info.get('birth_date', '待补充')}
联系电话：{employee_info.get('phone', '待补充')}
电子邮箱：{employee_info.get('email', '待补充')}
政治面貌：{employee_info.get('political_status', '待补充')}

【二、住址信息】
户籍地址：{employee_info.get('permanent_address', '待补充')}
现住址：{employee_info.get('current_address', '待补充')}

【三、教育信息】
最高学历：{employee_info.get('education', '待补充')}
毕业院校：{employee_info.get('school', '待补充')}
专业：{employee_info.get('major', '待补充')}
毕业时间：{employee_info.get('graduation_date', '待补充')}
学习方式：{employee_info.get('mode_of_study', '待补充')}

【四、任职信息】
部门：{employee_info.get('department', '待补充')}
职位：{employee_info.get('position', '待补充')}
入职日期：{employee_info.get('hire_date', '待补充')}
职级：{employee_info.get('job_level', '待补充')}
合同类型：{employee_info.get('contract_type', '待补充')}
合同期限：{employee_info.get('contract_duration', '待补充')}

{'='*40}
档案生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
{'='*40}
"""
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return True
        except Exception as e:
            print(f"生成文本档案失败: {e}")
            return False


# 全局实例
archive_generator = ArchiveGenerator()